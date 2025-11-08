# ============================================================
#  Kisbye Consulting – BusinessCaseGPT
#  Version: v9.1 (samlet, med logo + PyInstaller-stier)
# ============================================================

import os
import sys
import io
import json
import time
import shutil
import threading
import webbrowser
from datetime import datetime

from flask import (
    Flask, request, render_template_string, send_from_directory,
    Response, jsonify, url_for
)

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# AFSNIT 0 – STIER (virker i .py og i PyInstaller .exe)
# ============================================================
if getattr(sys, "frozen", False):
    # kører som exe
    BASE_DIR = sys._MEIPASS
    RUN_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    RUN_DIR = BASE_DIR

script_dir = RUN_DIR
OUTPUT_DIR = os.path.join(RUN_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

APP_TITLE = "Kisbye Consulting – BusinessCaseGPT v9.1"
LOGO_PNG_SOURCE = "kisbye_logo.png"
LOGO_ICO_SOURCE = "kisbye_logo.ico"

app = Flask(__name__, static_folder="static")

last_ping = time.time()  # til idle-killer


# ============================================================
# AFSNIT 1 – HJÆLPERE
# ============================================================
def ensure_output_dir() -> str:
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR, exist_ok=True)
    return OUTPUT_DIR


def safe_name(s: str) -> str:
    s = (s or "").strip()
    s = s.replace("æ", "ae").replace("Æ", "Ae")
    s = s.replace("ø", "oe").replace("Ø", "Oe")
    s = s.replace("å", "aa").replace("Å", "Aa")
    for bad in [" ", "/", "\\", ":", "*", "?", "\"", "<", ">", "|"]:
        s = s.replace(bad, "_")
    return s


def to_number(v, default=0.0) -> float:
    if v is None:
        return default
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if not s:
        return default
    s = s.replace(" ", "")
    s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return default


def fmt_num(x, decimals=0):
    if isinstance(x, str):
        x = to_number(x, 0.0)
    s = f"{x:.{decimals}f}"
    s = s.replace(".", ",")
    return s


def fmt_dkk(x, decimals=0) -> str:
    return f"{fmt_num(x, decimals)} kr"


def get_logo_path_for_docs():
    """
    Find kisbye_logo.png både når vi kører som .py og som PyInstaller .exe.
    Vi tjekker:
    1) static/ ved siden af script/EXE
    2) samme mappe som script/EXE
    3) static/ inde i PyInstallers _MEIPASS
    4) roden af _MEIPASS
    """
    # 1) ved siden af EXE / .py
    cand1 = os.path.join(script_dir, "static", "kisbye_logo.png")
    if os.path.exists(cand1):
        return cand1

    cand2 = os.path.join(script_dir, "kisbye_logo.png")
    if os.path.exists(cand2):
        return cand2

    # 2) hvis vi kører som frosset exe, så kig i _MEIPASS
    if getattr(sys, "frozen", False):
        base = sys._MEIPASS
        cand3 = os.path.join(base, "static", "kisbye_logo.png")
        if os.path.exists(cand3):
            return cand3
        cand4 = os.path.join(base, "kisbye_logo.png")
        if os.path.exists(cand4):
            return cand4

    # ellers ikke noget logo
    return None


def add_logo_header(doc: Document):
    """læg logo i header til venstre – crasher ikke hvis der mangler logo"""
    logo_path = get_logo_path_for_docs()
    if not logo_path:
        return
    try:
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.3))
    except Exception:
        pass


# ============================================================
# AFSNIT 2 – STANDARD-FORMULAR
# ============================================================
def empty_form() -> dict:
    return {
        "procesnavn": "",
        "formaal": "",
        "udfoerende": "",
        "proces_ejer": "",
        "sme": "",
        "rpa_udvikler": "",
        "sponsor": "",
        "systemer": "Excel, Outlook, SharePoint, Power Automate",
        "as_is_beskrivelse": "",
        "to_be_beskrivelse": "",
        "varighed_min": "35",
        "frekvens_pr_uge": "3",
        "arbejdsdage_pr_aar": "250",
        "aarSloen_kr": "450000",
        "automationsgrad_pct": "80",
        "investering_kr": "60000",
        "drift_aarlig_kr": "0",
        "kritikalitet": "Middel",
        "input": "Mail fra teamleder, Excel med medarbejderdata",
        "output": "Beregnet regneark, statusmail, logfil",
        "fejl": "Manglende data, forkert systemvalg, dobbeltindtastning",
        "kvalitative": "Færre fejl, hurtigere levering, bedre kvalitet",
        "rst_regel": "4",
        "rst_stabil": "3",
        "rst_tid": "3",
        "afhaengigheder": "Afhænger af HR-data, licens, godkendelse fra IT",
        "extra_json": "",
    }


# ============================================================
# AFSNIT 3 – BEREGNING
# ============================================================
def calc_metrics(c: dict) -> dict:
    varighed_min = to_number(c.get("varighed_min"), 0.0)
    frekvens_pr_uge = to_number(c.get("frekvens_pr_uge"), 0.0)
    aarsloen_kr = to_number(c.get("aarSloen_kr"), 450000.0)
    automationsgrad_pct = to_number(c.get("automationsgrad_pct"), 80.0)
    investering_kr = to_number(c.get("investering_kr"), 60000.0)
    drift_aarlig_kr = to_number(c.get("drift_aarlig_kr"), 0.0)

    minutter_pr_aar = varighed_min * frekvens_pr_uge * 52
    timer_pr_aar = minutter_pr_aar / 60.0
    fte = timer_pr_aar / 1540.0 if timer_pr_aar > 0 else 0.0
    timeloen = aarsloen_kr / 1540.0 if aarsloen_kr > 0 else 0.0

    omkostning_foer = timer_pr_aar * timeloen
    efter_timer = timer_pr_aar * (1 - automationsgrad_pct / 100.0)
    omkostning_efter = efter_timer * timeloen + drift_aarlig_kr

    aarlig_besparelse = omkostning_foer - omkostning_efter
    break_even_aar = investering_kr / aarlig_besparelse if aarlig_besparelse > 0 else 0

    return {
        "minutter_pr_aar": minutter_pr_aar,
        "timer_pr_aar": timer_pr_aar,
        "fte": fte,
        "timeloen": timeloen,
        "omkostning_foer": omkostning_foer,
        "omkostning_efter": omkostning_efter,
        "aarlig_besparelse": aarlig_besparelse,
        "break_even_aar": break_even_aar,
    }


# ============================================================
# AFSNIT 4 – WORD SPØRGESKEMA
# ============================================================
def build_word_questionnaire() -> bytes:
    doc = Document()

    # logo i header
    add_logo_header(doc)

    doc.add_heading("Business Case – spørgeskema", level=1)
    doc.add_paragraph("Udfyld felterne og upload dokumentet i BusinessCaseGPT.")

    questions = [
        ("Procesnavn", "procesnavn"),
        ("Formål", "formaal"),
        ("Udførende (roller/navne)", "udfoerende"),
        ("Procesejer", "proces_ejer"),
        ("Sponsor / bestiller", "sponsor"),
        ("SME / procesekspert", "sme"),
        ("RPA-udvikler", "rpa_udvikler"),
        ("Systemer i brug", "systemer"),
        ("Varighed pr. opgave (min)", "varighed_min"),
        ("Frekvens (gange/uge)", "frekvens_pr_uge"),
        ("Arbejdsdage pr. år", "arbejdsdage_pr_aar"),
        ("Årsløn (kr)", "aarSloen_kr"),
        ("Automatiseringsgrad (%)", "automationsgrad_pct"),
        ("Investering (kr)", "investering_kr"),
        ("Årlig licens/drift (kr)", "drift_aarlig_kr"),
        ("Input", "input"),
        ("Output", "output"),
        ("Typiske fejl/undtagelser", "fejl"),
        ("Kvalitative gevinster", "kvalitative"),
        ("AS-IS beskrivelse (sådan gør vi i dag)", "as_is_beskrivelse"),
        ("TO-BE beskrivelse (sådan skal robotten gøre)", "to_be_beskrivelse"),
        ("Afhængigheder", "afhaengigheder"),
    ]
    for title, _key in questions:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}: ")
        r.bold = True
        p.add_run("")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def parse_docx_to_form(file_storage) -> dict:
    try:
        doc = Document(file_storage)
    except Exception:
        return None

    mapping = empty_form()
    for para in doc.paragraphs:
        line = (para.text or "").strip()
        if not line or ":" not in line:
            continue
        keypart, val = line.split(":", 1)
        key = keypart.lower().strip()
        val = val.strip()

        if "procesnavn" in key:
            mapping["procesnavn"] = val
        elif "formål" in key or "formaal" in key:
            mapping["formaal"] = val
        elif "udførende" in key or "udfoerende" in key:
            mapping["udfoerende"] = val
        elif "procesejer" in key or "proces-ejer" in key:
            mapping["proces_ejer"] = val
        elif "sponsor" in key:
            mapping["sponsor"] = val
        elif "sme" in key:
            mapping["sme"] = val
        elif "rpa" in key:
            mapping["rpa_udvikler"] = val
        elif "systemer" in key:
            mapping["systemer"] = val
        elif "varighed" in key:
            mapping["varighed_min"] = val
        elif "frekvens" in key:
            mapping["frekvens_pr_uge"] = val
        elif "arbejdsdage" in key:
            mapping["arbejdsdage_pr_aar"] = val
        elif "årsløn" in key or "aarsløn" in key or "årsloen" in key:
            mapping["aarSloen_kr"] = val
        elif "automatiseringsgrad" in key:
            mapping["automationsgrad_pct"] = val
        elif "investering" in key:
            mapping["investering_kr"] = val
        elif "licens" in key or "drift" in key:
            mapping["drift_aarlig_kr"] = val
        elif key.startswith("input"):
            mapping["input"] = val
        elif key.startswith("output"):
            mapping["output"] = val
        elif "fejl" in key or "undtagelser" in key:
            mapping["fejl"] = val
        elif "kvalitative" in key:
            mapping["kvalitative"] = val
        elif "as-is" in key:
            mapping["as_is_beskrivelse"] = val
        elif "to-be" in key:
            mapping["to_be_beskrivelse"] = val
        elif "afhæng" in key or "afhaeng" in key:
            mapping["afhaengigheder"] = val

    return mapping


# ============================================================
# AFSNIT 5 – WORD PDD / RTS
# ============================================================
def build_word_pdd(path: str, c: dict, m: dict):
    doc = Document()

    # logo i header
    add_logo_header(doc)

    doc.add_heading(f"PDD / RTS – {c.get('procesnavn','Proces')}", level=1)

    # Overblik
    doc.add_heading("Overblik", level=2)
    doc.add_paragraph(f"Område: {c.get('formaal','HR / IT / Forretning')}")
    doc.add_paragraph(f"Procesejer: {c.get('proces_ejer','')}")
    doc.add_paragraph(f"Sponsor: {c.get('sponsor','')}")
    doc.add_paragraph(f"Udførende i dag: {c.get('udfoerende','')}")
    doc.add_paragraph(f"Systemer: {c.get('systemer','')}")

    # Formål
    doc.add_heading("Formål", level=2)
    doc.add_paragraph(
        "At dokumentere den nuværende (AS-IS) proces og beskrive den fremtidige (TO-BE) automatiserede proces, "
        "så RPA-udvikleren kan bygge, og ledelsen kan godkende."
    )

    # Interessenter
    doc.add_heading("Interessenter", level=2)
    doc.add_paragraph(f"- Procesejer / godkender: {c.get('proces_ejer','')}")
    doc.add_paragraph(f"- SME / procesekspert: {c.get('sme','')}")
    doc.add_paragraph(f"- RPA-udvikler: {c.get('rpa_udvikler','')}")
    doc.add_paragraph(f"- Sponsor / ledelse: {c.get('sponsor','')}")

    # AS-IS
    doc.add_heading("AS-IS proces", level=2)
    doc.add_paragraph(c.get("as_is_beskrivelse", "Manuel proces med flere aktører."))

    # TO-BE
    doc.add_heading("TO-BE proces (RPA / PAD)", level=2)
    doc.add_paragraph(
        c.get("to_be_beskrivelse", "Proces automatiseres, robotten henter input, opretter i systemer og logger resultat.")
    )

    # Input / Output / Fejl
    doc.add_heading("Input", level=3)
    doc.add_paragraph(c.get("input", ""))
    doc.add_heading("Output", level=3)
    doc.add_paragraph(c.get("output", ""))
    doc.add_heading("Fejl / undtagelser", level=3)
    doc.add_paragraph(c.get("fejl", ""))

    # Økonomi
    doc.add_heading("Økonomi (nøgletal)", level=2)
    doc.add_paragraph(f"Årligt tidsforbrug før automation: {fmt_num(m['timer_pr_aar'], 1)} timer")
    doc.add_paragraph(f"Årlig besparelse: {fmt_dkk(m['aarlig_besparelse'], 0)}")
    doc.add_paragraph(f"Investering: {fmt_dkk(c.get('investering_kr'), 0)}")
    doc.add_paragraph(f"Break-even: {fmt_num(m['break_even_aar'], 1)} år")

    # lidt større skrift
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)

    doc.save(path)


# ============================================================
# AFSNIT 6 – WORD LEDELSESBESKRIVELSE
# ============================================================
def build_word_leadership(path: str, c: dict, m: dict, extra_json_text: str = ""):
    doc = Document()

    # logo i header
    add_logo_header(doc)

    # Titel
    title = doc.add_heading(f"Ledelsesbeskrivelse – {c.get('procesnavn','Proces')}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph(
        "Formålet med dette dokument er at give ledelsen et klart beslutningsgrundlag for at automatisere processen."
    )
    intro.runs[0].font.size = Pt(11)

    # Executive summary som tabel
    doc.add_paragraph("")
    doc.add_paragraph("Executive summary:").runs[0].bold = True

    table = doc.add_table(rows=0, cols=2)
    rows = [
        ("Problem / nuværende situation", c.get("formaal", "Manuel proces med spildtid og fejl.")),
        ("Løsning", f"Automatiseret RPA/PAD-flow i {c.get('systemer','relevante systemer')}"),
        ("Tidsforbrug før", f"{c.get('varighed_min','?')} min × {c.get('frekvens_pr_uge','?')} pr. uge"),
        ("Automationsgrad", f"{c.get('automationsgrad_pct','80')} %"),
        ("Årlig besparelse", fmt_dkk(m["aarlig_besparelse"], 0)),
        ("Investering", fmt_dkk(c.get("investering_kr"), 0)),
        ("Break-even", f"{fmt_num(m['break_even_aar'], 1)} år"),
        ("Kvalitative gevinster", c.get("kvalitative", "Færre fejl, hurtigere levering, bedre service")),
    ]
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    # 1. Baggrund
    doc.add_heading("1. Baggrund og formål", level=2)
    doc.add_paragraph(
        "Processen udføres i dag manuelt af én eller flere roller. Det giver risiko for manglende data, dobbeltindtastning og ventetid. "
        "Automatiseringen skal standardisere opgaven og frigive tid til andre opgaver."
    )

    # 2. AS-IS / TO-BE
    doc.add_heading("2. Procesbeskrivelse (AS-IS → TO-BE)", level=2)
    doc.add_paragraph("AS-IS:").runs[0].bold = True
    doc.add_paragraph(c.get("as_is_beskrivelse", "Manuel proces uden standardisering."))
    doc.add_paragraph("TO-BE:").runs[0].bold = True
    doc.add_paragraph(c.get("to_be_beskrivelse", "Proces køres som RPA-flow/PAD med faste input og logning."))

    # 3. Økonomi
    doc.add_heading("3. Økonomi", level=2)
    doc.add_paragraph(f"Årligt tidsforbrug før automation: {fmt_num(m['timer_pr_aar'], 1)} timer.")
    doc.add_paragraph(f"Årlig omkostning før: {fmt_dkk(m['omkostning_foer'], 0)}.")
    doc.add_paragraph(f"Årlig omkostning efter: {fmt_dkk(m['omkostning_efter'], 0)}.")
    doc.add_paragraph(f"Forventet årlig besparelse: {fmt_dkk(m['aarlig_besparelse'], 0)}.")
    doc.add_paragraph(f"Investering: {fmt_dkk(c.get('investering_kr'), 0)}.")
    doc.add_paragraph(f"Break-even: {fmt_num(m['break_even_aar'], 1)} år.")

    # 4. Roller
    doc.add_heading("4. Roller og ansvar", level=2)
    doc.add_paragraph(f"Procesejer: {c.get('proces_ejer','')}")
    doc.add_paragraph(f"Sponsor: {c.get('sponsor','')}")
    doc.add_paragraph(f"SME / procesekspert: {c.get('sme','')}")
    doc.add_paragraph(f"RPA-udvikler: {c.get('rpa_udvikler','')}")

    # 5. Gevinster
    doc.add_heading("5. Gevinster (kvalitative)", level=2)
    doc.add_paragraph(
        c.get(
            "kvalitative",
            "Hurtigere levering, færre fejl, bedre datakvalitet, tilfredse medarbejdere."
        )
    )

    # 6. Risiko
    doc.add_heading("6. Risiko og afhængigheder", level=2)
    doc.add_paragraph(c.get("afhaengigheder", "Afhænger af adgang til HR-/fagsystemer og licenser."))

    # 7. Konklusion
    doc.add_heading("7. Konklusion og anbefaling", level=2)
    doc.add_paragraph(
        "Automatiseringen kan gennemføres med lav til middel risiko og med tydelig økonomisk effekt. "
        "Det anbefales, at ledelsen godkender projektet og igangsætter udviklingen."
    )

    if extra_json_text:
        doc.add_heading("Bilag – rådata fra formular/upload", level=2)
        doc.add_paragraph(extra_json_text)

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)

    doc.save(path)


# ============================================================
# AFSNIT 7 – EXCEL
# ============================================================
YELLOW = PatternFill("solid", fgColor="FFF2CC")
GREY = PatternFill("solid", fgColor="F2F2F2")

def build_excel(path: str, c: dict, m: dict):
    from openpyxl.drawing.image import Image as XLImage

    wb = Workbook()
    ws = wb.active
    ws.title = "Forside"
    ws["A1"] = "RPA Business Case"
    ws["A1"].font = Font(size=16, bold=True)
    ws["A3"] = "Procesnavn"
    ws["B3"] = c.get("procesnavn", "")
    ws["A4"] = "Formål"
    ws["B4"] = c.get("formaal", "")
    ws["A5"] = "Procesejer"
    ws["B5"] = c.get("proces_ejer", "")
    ws["A6"] = "Dato"
    ws["B6"] = datetime.now().strftime("%d-%m-%Y")

    # logo i excel hvis muligt
    logo_path = get_logo_path_for_docs()
    if logo_path and os.path.exists(logo_path):
        try:
            img = XLImage(logo_path)
            img.height = 90
            img.width = 180
            ws.add_image(img, "D1")
        except Exception:
            pass

    # Spørgsmål
    ws2 = wb.create_sheet("Spørgsmål")
    ws2["A1"] = "Felt"
    ws2["B1"] = "Værdi"
    row = 2
    for key, val in c.items():
        ws2[f"A{row}"] = key
        ws2[f"B{row}"] = val
        row += 1
    ws2.column_dimensions["A"].width = 35
    ws2.column_dimensions["B"].width = 80

    # PDD
    ws3 = wb.create_sheet("PDD (AS-IS)")
    ws3["A1"] = "AS-IS beskrivelse"
    ws3["A2"] = c.get("as_is_beskrivelse", "")
    ws3["A2"].alignment = Alignment(wrap_text=True)

    # RTS
    ws4 = wb.create_sheet("RTS (TO-BE)")
    ws4["A1"] = "TO-BE beskrivelse"
    ws4["A2"] = c.get("to_be_beskrivelse", "")
    ws4["A2"].alignment = Alignment(wrap_text=True)

    # Økonomi
    ws5 = wb.create_sheet("Økonomi")
    ws5["A1"] = "Parameter"
    ws5["B1"] = "Værdi"
    ws5["A2"] = "Minutter pr. år"
    ws5["B2"] = m["minutter_pr_aar"]
    ws5["A3"] = "Timer pr. år"
    ws5["B3"] = m["timer_pr_aar"]
    ws5["A4"] = "Årlig omkostning før"
    ws5["B4"] = m["omkostning_foer"]
    ws5["A5"] = "Årlig omkostning efter"
    ws5["B5"] = m["omkostning_efter"]
    ws5["A6"] = "Årlig besparelse"
    ws5["B6"] = m["aarlig_besparelse"]
    ws5["A7"] = "Investering"
    ws5["B7"] = to_number(c.get("investering_kr"), 0.0)
    ws5["A8"] = "Break-even (år)"
    ws5["B8"] = m["break_even_aar"]

    # Business Case
    ws6 = wb.create_sheet("Business Case")
    ws6["A1"] = "Business Case – samlet vurdering"
    ws6["A3"] = "Anbefaling"
    ws6["B3"] = "Automatisering anbefales – lav risiko, hurtig gevinst."
    ws6["A5"] = "Kvalitative gevinster"
    ws6["B5"] = c.get("kvalitative", "")

    # Ledelse – kort overblik
    ws7 = wb.create_sheet("Ledelse")
    ws7["A1"] = "Ledelsesoverblik"
    ws7["A1"].font = Font(size=14, bold=True)
    ws7["A3"] = "Procesnavn"
    ws7["B3"] = c.get("procesnavn", "")
    ws7["A4"] = "Årlig besparelse"
    ws7["B4"] = m["aarlig_besparelse"]
    ws7["A5"] = "Investering"
    ws7["B5"] = to_number(c.get("investering_kr"), 0.0)
    ws7["A6"] = "Break-even (år)"
    ws7["B6"] = m["break_even_aar"]
    ws7["A8"] = "Anbefaling"
    ws7["B8"] = "Automatisering anbefales – lav risiko, tydelig effekt."
    ws7.column_dimensions["A"].width = 28
    ws7.column_dimensions["B"].width = 50

    wb.save(path)


# ============================================================
# AFSNIT 8 – HTML TEMPLATES
# ============================================================
# ============================================================
# AFSNIT – HTML TEMPLATES (forside + resultat)
# ============================================================
FORM_HTML = r"""
<!doctype html>
<html lang="da">
<head>
  <meta charset="utf-8">
  <title>{{ title }}</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
  <style>
    body { background:#f4f6fb; }
    .card { border-radius:16px; box-shadow:0 8px 30px rgba(0,0,0,.06); }
    .brand { display:flex; flex-direction:column; align-items:center; gap:.5rem; }
    .brand img.logo { max-height:110px; width:auto; }
    .section-title { font-weight:700; color:#334; margin-top:22px; }
  </style>
</head>
<body>
<div class="container py-4">
  <!-- TOP / LOGO -->
  <div class="brand mb-3">
    {% if logo_png %}
      <img src="/static/kisbye_logo.png" alt="Kisbye Consulting" class="logo">
    {% elif logo_ico %}
      <img src="/static/kisbye_logo.ico" alt="Kisbye Consulting" width="84" height="84">
    {% endif %}
    <h3 class="text-center">{{ title }}</h3>
    <p class="text-muted mb-0">Indtast én proces – få Excel + 2 Word-dokumenter</p>
  </div>

  <!-- KORT MED FORMULAR -->
  <div class="card p-4 mb-4">
    <!-- VÆRKTØJSLINJE -->
    <div class="d-flex flex-wrap gap-2 justify-content-between align-items-center mb-3">
      <h4 class="mb-0">1. Grundlæggende oplysninger</h4>
      <div class="d-flex gap-2">
        <a href="{{ url_for('download_word_template') }}" class="btn btn-outline-primary btn-sm">
          Download spørgeskema (Word)
        </a>
        <form action="{{ url_for('load_json') }}" method="post" enctype="multipart/form-data" class="d-flex gap-2">
          <input type="file" name="jsonfile" accept=".json,.txt" class="form-control form-control-sm">
          <button class="btn btn-outline-secondary btn-sm" type="submit">Udfyld fra JSON</button>
        </form>
        <form action="{{ url_for('load_docx') }}" method="post" enctype="multipart/form-data" class="d-flex gap-2">
          <input type="file" name="docxfile" accept=".docx" class="form-control form-control-sm">
          <button class="btn btn-outline-secondary btn-sm" type="submit">Udfyld fra Word</button>
        </form>
      </div>
    </div>
    <p class="text-muted mb-3">Upload en tidligere JSON eller Word – eller udfyld felterne nedenfor.</p>

    <!-- FORMULAR START -->
    <form method="post" action="{{ url_for('generate') }}">
      <div class="row g-3">
        <div class="col-md-6">
          <label class="form-label">Procesnavn</label>
          <input name="procesnavn" class="form-control" value="{{ f.procesnavn }}" required>
        </div>
        <div class="col-md-6">
          <label class="form-label">Formål</label>
          <input name="formaal" class="form-control" value="{{ f.formaal }}" required>
        </div>

        <div class="col-md-6">
          <label class="form-label">Udførende (dem der gør det i dag)</label>
          <input name="udfoerende" class="form-control" value="{{ f.udfoerende }}">
        </div>
        <div class="col-md-6">
          <label class="form-label">Proces-ejer / forretningsansvarlig</label>
          <input name="proces_ejer" class="form-control" value="{{ f.proces_ejer }}">
        </div>
        <div class="col-md-6">
          <label class="form-label">SME / fagperson</label>
          <input name="sme" class="form-control" value="{{ f.sme }}">
        </div>
        <div class="col-md-6">
          <label class="form-label">RPA-udvikler / teknisk ansvarlig</label>
          <input name="rpa_udvikler" class="form-control" value="{{ f.rpa_udvikler }}">
        </div>
        <div class="col-md-6">
          <label class="form-label">Sponsor / godkender</label>
          <input name="sponsor" class="form-control" value="{{ f.sponsor }}">
        </div>
        <div class="col-md-6">
          <label class="form-label">Systemer i brug</label>
          <input name="systemer" class="form-control" value="{{ f.systemer }}">
        </div>
      </div>

      <h4 class="section-title">2. Tidsforbrug og volumener</h4>
      <div class="row g-3">
        <div class="col-md-3">
          <label class="form-label">Varighed pr. opgave (min)</label>
          <input type="number" name="varighed_min" class="form-control" value="{{ f.varighed_min }}" required>
        </div>
        <div class="col-md-3">
          <label class="form-label">Frekvens (gange/uge)</label>
          <input type="number" step="0.1" name="frekvens_pr_uge" class="form-control" value="{{ f.frekvens_pr_uge }}" required>
        </div>
        <div class="col-md-3">
          <label class="form-label">Arbejdsdage pr. år</label>
          <input type="number" name="arbejdsdage_pr_aar" class="form-control" value="{{ f.arbejdsdage_pr_aar }}" required>
        </div>
        <div class="col-md-3">
          <label class="form-label">Årsløn (kr)</label>
          <input type="number" name="aarSloen_kr" class="form-control" value="{{ f.aarSloen_kr }}" required>
        </div>
      </div>

      <h4 class="section-title">3. Økonomi og investering</h4>
      <div class="row g-3">
        <div class="col-md-3">
          <label class="form-label">Automationsgrad (%)</label>
          <input type="number" name="automationsgrad_pct" class="form-control" value="{{ f.automationsgrad_pct }}" required>
        </div>
        <div class="col-md-3">
          <label class="form-label">Investering (kr)</label>
          <input type="number" name="investering_kr" class="form-control" value="{{ f.investering_kr }}" required>
        </div>
        <div class="col-md-3">
          <label class="form-label">Årlig drift/licens (kr)</label>
          <input type="number" name="drift_aarlig_kr" class="form-control" value="{{ f.drift_aarlig_kr }}">
        </div>
        <div class="col-md-3">
          <label class="form-label">Proceskritikalitet</label>
          <select name="kritikalitet" class="form-select">
            <option value="Høj" {% if f.kritikalitet=='Høj' %}selected{% endif %}>Høj</option>
            <option value="Middel" {% if f.kritikalitet=='Middel' %}selected{% endif %}>Middel</option>
            <option value="Lav" {% if f.kritikalitet=='Lav' %}selected{% endif %}>Lav</option>
          </select>
        </div>
      </div>

      <h4 class="section-title">4. Fejl, input, output</h4>
      <div class="row g-3">
        <div class="col-md-6">
          <label class="form-label">Input</label>
          <input name="input" class="form-control" value="{{ f.input }}">
        </div>
        <div class="col-md-6">
          <label class="form-label">Output</label>
          <input name="output" class="form-control" value="{{ f.output }}">
        </div>
        <div class="col-md-12">
          <label class="form-label">Typiske fejl/undtagelser</label>
          <input name="fejl" class="form-control" value="{{ f.fejl }}">
        </div>
        <div class="col-md-12">
          <label class="form-label">Kvalitative gevinster</label>
          <input name="kvalitative" class="form-control" value="{{ f.kvalitative }}">
        </div>
      </div>

      <h4 class="section-title">5. AS-IS og TO-BE</h4>
      <div class="row g-3">
        <div class="col-md-12">
          <label class="form-label">AS-IS – hvordan gør I i dag?</label>
          <textarea name="as_is_beskrivelse" class="form-control" rows="4">{{ f.as_is_beskrivelse }}</textarea>
        </div>
        <div class="col-md-12">
          <label class="form-label">TO-BE – hvordan skal robotten gøre?</label>
          <textarea name="to_be_beskrivelse" class="form-control" rows="3">{{ f.to_be_beskrivelse }}</textarea>
        </div>
      </div>

      <h4 class="section-title">6. Ekstra JSON / rå data</h4>
      <div class="row g-3">
        <div class="col-md-12">
          <textarea name="extra_json" class="form-control" rows="4">{{ f.extra_json }}</textarea>
        </div>
      </div>

      <div class="mt-4 d-flex gap-2">
        <button class="btn btn-primary btn-lg" type="submit">Generér Business Case</button>
        <a class="btn btn-outline-secondary" href="{{ url_for('index') }}">Nulstil</a>
      </div>
    </form>
  </div>
</div>
</body>
</html>
"""

RESULT_HTML = r"""
<!doctype html>
<html lang="da">
<head>
  <meta charset="utf-8">
  <title>Resultat – BusinessCaseGPT v9.1</title>
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
</head>
<body class="bg-light">
<div class="container py-4">
  <div class="card p-4">
    <h3 class="mb-2">✅ Business case genereret</h3>
    <p class="text-muted">Filerne er gemt i <code>{{ outdir }}</code></p>
    <ul>
      <li>📊 <a href="{{ excel_url }}">Excel – Business Case</a></li>
      <li>📝 <a href="{{ pdd_url }}">Word – PDD + RTS</a></li>
      <li>📋 <a href="{{ lb_url }}">Word – Ledelsesbeskrivelse</a></li>
    </ul>
    <div class="d-flex gap-2 mt-3">
      <a href="{{ url_for('index') }}" class="btn btn-primary">Ny Business Case</a>
      <button id="exitBtn" class="btn btn-outline-danger">Afslut program</button>
    </div>
  </div>
</div>
<script>
document.getElementById("exitBtn").addEventListener("click", async ()=>{
  try { await fetch("/shutdown", {method:"POST"}); } catch(e){}
  window.close();
});
</script>
</body>
</html>
"""


# ============================================================
# AFSNIT 9 – ROUTES
# ============================================================
@app.after_request
def no_cache(resp):
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    return resp


@app.route("/", methods=["GET"])
def index():
    global last_ping
    last_ping = time.time()
    return render_template_string(
        FORM_HTML,
        title=APP_TITLE,
        logo_png=os.path.exists(os.path.join("static", "kisbye_logo.png")),
        logo_ico=os.path.exists(os.path.join("static", "kisbye_logo.ico")),
        f=empty_form(),
    )


@app.route("/download_word_template", methods=["GET"])
def download_word_template():
    content = build_word_questionnaire()
    return Response(
        content,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=businesscase_spoergeskema.docx"},
    )


@app.route("/load_json", methods=["POST"])
def load_json():
    global last_ping
    last_ping = time.time()

    f = empty_form()
    file = request.files.get("jsonfile")
    if not file:
        return render_template_string(
            FORM_HTML,
            title=APP_TITLE,
            logo_png=os.path.exists(os.path.join("static", "kisbye_logo.png")),
            logo_ico=os.path.exists(os.path.join("static", "kisbye_logo.ico")),
            f=f,
        )
    try:
        data = json.loads(file.read().decode("utf-8"))
    except Exception as e:
        f["extra_json"] = f"Kunne ikke læse JSON: {e}"
        return render_template_string(
            FORM_HTML,
            title=APP_TITLE,
            logo_png=os.path.exists(os.path.join("static", "kisbye_logo.png")),
            logo_ico=os.path.exists(os.path.join("static", "kisbye_logo.ico")),
            f=f,
        )

    for key in f.keys():
        if key in data and not isinstance(data[key], dict):
            f[key] = str(data[key])

    po = data.get("process_overview") or {}
    if po:
        f["procesnavn"] = po.get("process_name", f["procesnavn"])
        f["formaal"] = po.get("objective", f["formaal"])
        systems = po.get("systems_in_scope")
        if isinstance(systems, list):
            f["systemer"] = ", ".join(systems)
        elif isinstance(systems, str):
            f["systemer"] = systems

    ta = data.get("timing_analysis") or {}
    if ta:
        if "minutes_per_hire" in ta:
            f["varighed_min"] = str(ta["minutes_per_hire"])
        workdays = po.get("workdays_per_year") or ta.get("workdays_per_year")
        if workdays:
            f["arbejdsdage_pr_aar"] = str(workdays)
        freq = ta.get("frequency_per_week")
        if freq:
            f["frekvens_pr_uge"] = str(freq)

    f["extra_json"] = json.dumps(data, indent=2, ensure_ascii=False)

    return render_template_string(
        FORM_HTML,
        title=APP_TITLE,
        logo_png=os.path.exists(os.path.join("static", "kisbye_logo.png")),
        logo_ico=os.path.exists(os.path.join("static", "kisbye_logo.ico")),
        f=f,
    )


@app.route("/load_docx", methods=["POST"])
def load_docx():
    global last_ping
    last_ping = time.time()

    file = request.files.get("docxfile")
    if not file:
        f = empty_form()
        f["extra_json"] = "Ingen Word-fil valgt."
        return render_template_string(
            FORM_HTML,
            title=APP_TITLE,
            logo_png=os.path.exists(os.path.join("static", "kisbye_logo.png")),
            logo_ico=os.path.exists(os.path.join("static", "kisbye_logo.ico")),
            f=f,
        )

    try:
        filled = parse_docx_to_form(file)
    except Exception:
        filled = None

    if filled is None:
        filled = empty_form()
        filled["extra_json"] = "Kunne ikke læse Word-filen – tjek formatet."

    return render_template_string(
        FORM_HTML,
        title=APP_TITLE,
        logo_png=os.path.exists(os.path.join("static", "kisbye_logo.png")),
        logo_ico=os.path.exists(os.path.join("static", "kisbye_logo.ico")),
        f=filled,
    )


@app.route("/generate", methods=["POST"])
def generate():
    global last_ping
    last_ping = time.time()

    c = empty_form()
    for key in c.keys():
        if key in request.form:
            c[key] = request.form.get(key, "").strip()

    m = calc_metrics(c)

    outdir = ensure_output_dir()
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    base = safe_name(c.get("procesnavn") or "RPA_BusinessCase")

    excel_path = os.path.join(outdir, f"{base}_BC_{stamp}.xlsx")
    pdd_path = os.path.join(outdir, f"{base}_PDD_RTS_{stamp}.docx")
    lb_path = os.path.join(outdir, f"{base}_Ledelsesbeskrivelse_{stamp}.docx")

    build_excel(excel_path, c, m)
    build_word_pdd(pdd_path, c, m)
    build_word_leadership(lb_path, c, m, extra_json_text=c.get("extra_json", ""))

    return render_template_string(
        RESULT_HTML,
        outdir=outdir,
        excel_url=f"/output/{os.path.basename(excel_path)}",
        pdd_url=f"/output/{os.path.basename(pdd_path)}",
        lb_url=f"/output/{os.path.basename(lb_path)}",
    )


@app.route("/output/<path:filename>")
def download_file(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


@app.route("/shutdown", methods=["POST"])
def shutdown():
    def delayed():
        time.sleep(0.3)
        os._exit(0)
    threading.Thread(target=delayed, daemon=True).start()
    return jsonify({"status": "ok"})


# ============================================================
# AFSNIT 10 – MAIN
# ============================================================
def idle_killer(seconds=180):
    global last_ping
    while True:
        time.sleep(10)
        try:
            if time.time() - last_ping > seconds:
                print("[idle-killer] Ingen aktivitet – lukker ned.")
                os._exit(0)
        except Exception:
            pass


def open_browser():
    try:
        time.sleep(0.7)
        webbrowser.open("http://127.0.0.1:5000", new=2)
    except Exception as e:
        print("[open_browser] Kunne ikke åbne browser:", e)


if __name__ == "__main__":
    # sørg for static-mappe
    os.makedirs(os.path.join(script_dir, "static"), exist_ok=True)

    # kopier logo hvis det ligger ved siden af
    src_png = os.path.join(script_dir, LOGO_PNG_SOURCE)
    dst_png = os.path.join(script_dir, "static", "kisbye_logo.png")
    if os.path.exists(src_png) and os.path.abspath(src_png) != os.path.abspath(dst_png):
        try:
            shutil.copy(src_png, dst_png)
        except Exception:
            pass

    src_ico = os.path.join(script_dir, LOGO_ICO_SOURCE)
    dst_ico = os.path.join(script_dir, "static", "kisbye_logo.ico")
    if os.path.exists(src_ico) and os.path.abspath(src_ico) != os.path.abspath(dst_ico):
        try:
            shutil.copy(src_ico, dst_ico)
        except Exception:
            pass

    threading.Thread(target=idle_killer, args=(180,), daemon=True).start()
    threading.Thread(target=open_browser, daemon=True).start()

    print("Kører på http://127.0.0.1:5000")
    app.run(host="127.0.0.1", port=5000, debug=False)
